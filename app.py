import os
import zipfile
from flask import Flask, render_template, request, redirect, flash, send_file, url_for, session, jsonify, g
from werkzeug.utils import secure_filename
from PIL import Image
from pdf2image import convert_from_path
from moviepy.editor import VideoFileClip, AudioFileClip
from pydub import AudioSegment
import pandas as pd
import docx
from docx import Document
import pdfplumber
import pypandoc
from pypandoc import convert_file as pypandoc_convert
from docx2pdf import convert as docx2pdf_convert
from pdf2docx import Converter
import pdfkit
import docx2txt
import csv
from PyPDF2 import PdfReader
import logging
import mysql.connector  # MySQL connector
from mysql.connector import Error  # Error handling
import time
import subprocess
import bcrypt
import datetime
from datetime import timedelta
#from flask import jsonify, request, render_template
import threading

app = Flask(__name__)
app.secret_key = os.urandom(24)

# Configuration
UPLOAD_FOLDER = 'uploads/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
logging.basicConfig(level=logging.DEBUG)
os.environ["PATH"] += os.pathsep + os.path.join(os.path.dirname(os.path.abspath(__file__)), 'bin')
pdf_config = pdfkit.configuration(wkhtmltopdf=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'bin', 'wkhtmltopdf.exe'))

# Ensure the upload folder exists
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'webp'}
ALLOWED_AUDIO_EXTENSIONS = {'mp3', 'wav', 'ogg', 'aac', 'flac', 'm4a'}
ALLOWED_VIDEO_EXTENSIONS = {'mp4', 'avi', 'mov', 'mkv', 'flv'}
ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'xlsx','xls', 'csv', 'docx', 'txt'}

# Set the path to ffmpeg
AudioSegment.converter = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'bin', 'ffmpeg.exe')

conv_time = 1

def allowed_file(filename, allowed_extensions):
    """Check if the file extension is allowed."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

@app.route('/')
def home():
    return render_template('index.html')

def cleanup_file(file_path, retries=3, delay=1):
    """Try to delete a file with retry logic and return success status."""
    for attempt in range(retries):
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f"Deleted file: {file_path}")
                return True  # Return success status
            return False  # File did not exist to begin with
        except PermissionError:
            flash(f'PermissionError: Could not delete file {file_path}. Attempt {attempt + 1} of {retries}.')
            time.sleep(delay)  # Wait before retrying
        except Exception as e:
            flash(f'Error deleting file {file_path}: {str(e)}')
            return False  # Exit on other exceptions
    return False  # Return failure status after retries

def insert_file_for_deletion(file_name):
    conn = create_connection()
    with conn.cursor() as cursor:
        cursor.execute("INSERT INTO files_to_delete (file_name) VALUES (%s)", (file_name,))
    conn.commit()
    conn.close()

# Function to delete scheduled files
def delete_scheduled_files():
    while True:
        now = datetime.datetime.now()
        next_run_time = now.replace(hour=0, minute=0, second=0, microsecond=0) + datetime.timedelta(days=1)
        wait_time = (next_run_time - now).total_seconds()
        
        # Wait until the next midnight
        logging.info(f"Waiting for {wait_time} seconds until midnight.")
        time.sleep(wait_time)

        # Proceed with file deletion at midnight
        logging.info("Starting scheduled file deletion at midnight.")
        conn = create_connection()
        try:
            with conn.cursor() as cursor:
                cursor.execute("SELECT COUNT(*) FROM files_to_delete")
                (file_count,) = cursor.fetchone()  # Get the count of files

                if file_count > 0:  # Only proceed if there are files to delete
                    cursor.execute("SELECT file_name FROM files_to_delete")
                    files = cursor.fetchall()

                    for (file_name,) in files:
                        try:
                            if os.path.exists(file_name):
                                os.remove(file_name)
                                logging.debug(f"Deleted scheduled file: {file_name}")
                                # Remove the entry from the database
                                cursor.execute("DELETE FROM files_to_delete WHERE file_name = %s", (file_name,))
                                conn.commit()
                            else:
                                logging.warning(f"File not found: {file_name}")
                        except PermissionError as e:
                            logging.error(f"PermissionError: Could not delete {file_name}. Reason: {e}")
                        except Exception as e:
                            logging.error(f"Unexpected error while deleting {file_name}: {e}")

        except Exception as e:
            logging.error(f"Database error: {e}")
        finally:
            conn.close()

# Start the background thread when the application starts
threading.Thread(target=delete_scheduled_files, daemon=True).start()

@app.route('/image_convert', methods=['GET', 'POST'])
def image_convert():
    """Handle image file upload and conversion."""
    if request.method == 'POST':
        file = request.files.get('file')
        if file and allowed_file(file.filename, ALLOWED_IMAGE_EXTENSIONS):
            # Get the secure filename and save the file
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Extract the file extension from the filename
            original_format = os.path.splitext(filename)[1].replace('.', '')
            target_format = request.form.get('format')
            output_image_file = os.path.splitext(file_path)[0] + f'.{target_format}'

            # Initialize variables for conversion logging
            original_file_size = os.path.getsize(file_path)
            converted_file_size = None
            conversion_time = None
            conversion_status = 'failed'  # Default status for failure
            file_id = None  # Initialize file_id
            conversion_type = f"{original_format} to {target_format}"  # Set your conversion type accordingly

            try:
                # Start measuring conversion time
                start_time = time.time()

                # Process image conversion (replace with your image processing library)
                with Image.open(file_path) as img:
                    img.save(output_image_file, format=target_format)

                # End measuring conversion time
                end_time = time.time()
                conversion_time = end_time - start_time  # Calculate time in seconds as float

                # Get file sizes for logging
                original_file_size = os.path.getsize(file_path)
                converted_file_size = os.path.getsize(output_image_file)

                # If everything goes well, mark conversion as successful
                conversion_status = 'successful'

                # Log the file metadata and retrieve the file_id
                file_id = log_img_file_metadata(filename, original_format, target_format, conversion_type, conversion_status)

                # Log the successful conversion
                log_img_conversion(conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status, file_id)

                # Send the converted file back to the user
                return send_file(output_image_file, as_attachment=True, mimetype='image/' + target_format)

            except Exception as e:
                # Handle conversion errors
                logging.error(f"Conversion error: {e}")
                flash(f'Error processing image file: {e}', 'error')

                # Log the failed conversion with the existing file sizes and conversion time
                if file_id is None:
                    # Log the metadata first to get the file_id if conversion failed
                    file_id = log_img_file_metadata(filename, original_format, target_format, conversion_type, conversion_status)

                # Log the failed conversion
                log_img_conversion(conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status, file_id)

                return redirect(request.url)
            
            finally:
                time.sleep(conv_time)
                # Ensure files are deleted after closing video
                #cleanup_file(file_path)  # Delete original file
                if output_image_file:
                    insert_file_for_deletion(output_image_file)
                if file_path:
                    insert_file_for_deletion(file_path)

                # # Ensure files are deleted after closing video
                # cleanup_file(file_path)  # Delete original file

                # time.sleep(1)  # Wait for 10 seconds before attempting to delete the converted file

                # # Attempt to delete the converted file
                # if os.path.exists(output_image_file):
                #     print(f"Attempting to delete converted file: {output_image_file}")
                #     deletion_result = cleanup_file(output_image_file)  # Attempt to delete converted file
                #     if deletion_result:
                #         print(f"Successfully deleted converted file: {output_image_file}")
                #     else:
                #         print(f"Failed to delete converted file: {output_image_file}")

                # else:
                #     print(f"Converted file does not exist: {output_image_file}")
                
        else:
            flash('Invalid file format. Please upload a valid image file.', 'error')

    return render_template('image_convert.html')

@app.route('/audio_convert', methods=['GET', 'POST'])
def audio_convert():
    """Handle audio file upload and conversion."""
    if request.method == 'POST':
        file = request.files.get('file')
        if file and allowed_file(file.filename, ALLOWED_AUDIO_EXTENSIONS):
            # Get the secure filename and save the file
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Extract the file extension from the filename
            original_format = os.path.splitext(filename)[1].replace('.', '')
            target_format = request.form.get('format')
            output_audio_file = os.path.splitext(file_path)[0] + f'.{target_format}'

            # Initialize variables for conversion logging
            original_file_size = None
            converted_file_size = None
            conversion_time = None
            conversion_status = 'failed'  # Default status for failure
            file_id = None  # Initialize file_id
            conversion_type = f"{original_format} - {target_format}"  # Set your conversion type accordingly

            try:
                # Start measuring conversion time
                start_time = time.time()

                # Process audio conversion
                audio = AudioSegment.from_file(file_path)
                audio.export(output_audio_file, format=target_format)

                # End measuring conversion time
                end_time = time.time()
                conversion_time = end_time - start_time  # Calculate time in seconds as float

                # Get file sizes for logging
                original_file_size = os.path.getsize(file_path)
                converted_file_size = os.path.getsize(output_audio_file)

                # If everything goes well, mark conversion as successful
                conversion_status = 'successful'

                # Log the file metadata and retrieve the file_id
                file_id = log_audio_file_metadata(filename, original_format, target_format, conversion_status, conversion_type)

                # Log the successful conversion
                log_audio_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type)

                # Send the converted file back to the user
                return send_file(output_audio_file, as_attachment=True, mimetype='audio/' + target_format)

            except Exception as e:
                # Handle conversion errors
                logging.error(f"Conversion error: {e}")
                flash(f'Error processing audio file: {e}', 'error')

                # Log the failed conversion with the existing file sizes and conversion time
                if file_id is None:
                    # Log the metadata first to get the file_id if conversion failed
                    file_id = log_audio_file_metadata(filename, original_format, target_format, conversion_status, conversion_type)

                # Log the failed conversion
                log_audio_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type)

                return redirect(request.url)
            
            finally:
                time.sleep(conv_time)
                #cleanup_file(file_path)
                if output_audio_file:
                    insert_file_for_deletion(output_audio_file)
                if file_path:
                    insert_file_for_deletion(file_path)

                # # Ensure files are deleted after closing video
                # cleanup_file(file_path)  # Delete original file

                # time.sleep(1)  # Wait for 10 seconds before attempting to delete the converted file

                # # Attempt to delete the converted file
                # if os.path.exists(output_audio_file):
                #     print(f"Attempting to delete converted file: {output_audio_file}")
                #     deletion_result = cleanup_file(output_audio_file)  # Attempt to delete converted file
                #     if deletion_result:
                #         print(f"Successfully deleted converted file: {output_audio_file}")
                #     else:
                #         print(f"Failed to delete converted file: {output_audio_file}")

                # else:
                #     print(f"Converted file does not exist: {output_audio_file}")
        else:
            flash('Invalid file format. Please upload a valid audio file.', 'error')

    return render_template('audio_convert.html')

@app.route('/video_convert', methods=['GET', 'POST'])
def video_convert():
    """Handle video file upload and conversion."""
    MAX_FILE_SIZE = 300 * 1024 * 1024  # 300 MB in bytes

    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)

        file = request.files['file']

        if file and allowed_file(file.filename, ALLOWED_VIDEO_EXTENSIONS):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Check file size
            original_file_size = os.path.getsize(file_path)
            if original_file_size > MAX_FILE_SIZE:
                flash('File exceeds the maximum size limit of 300 MB.')
                os.remove(file_path)  # Optionally, remove the oversized file
                return redirect(request.url)

            target_format = request.form.get('format')
            output_file = os.path.splitext(file_path)[0] + '.' + target_format

            # Initialize conversion details
            conversion_status = 'failed'
            conversion_time = 0.0
            conversion_type = f"{file.content_type} - {target_format}"  # Set your conversion type accordingly

            try:
                # Start conversion and log time
                start_time = time.time()  # Use time to track conversion duration
                video = VideoFileClip(file_path)
                video.write_videofile(output_file, codec='libx264', bitrate='5000k', preset='slow', ffmpeg_params=['-crf', '18'])

                # Mark conversion as successful
                conversion_status = 'successful'
                conversion_time = time.time() - start_time  # Calculate conversion time in seconds
                converted_file_size = os.path.getsize(output_file)

                # Log conversion details to the database
                file_id = log_vid_file_metadata(filename, file.content_type, target_format, conversion_status, conversion_type)
                log_vid_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type)

                return send_file(output_file, as_attachment=True)

            except Exception as e:
                flash(f'Error processing file: {e}')
                # Log failed conversion
                log_vid_conversion(None, original_file_size, None, conversion_time, conversion_status, conversion_type)
                return redirect(request.url)
            
            finally:
                time.sleep(conv_time)
                #cleanup_file(file_path)
                if output_file:
                    insert_file_for_deletion(output_file)
                if file_path:
                    insert_file_for_deletion(file_path)
                # # Clean up temporary files
                # if video is not None:
                #     video.reader.close()
                #     if video.audio is not None:
                #         video.audio.reader.close_proc()

                # # Ensure files are deleted after closing video
                # cleanup_file(file_path)  # Delete original file

                # time.sleep(1)  # Wait for 10 seconds before attempting to delete the converted file

                # # Attempt to delete the converted file
                # if os.path.exists(output_file):
                #     print(f"Attempting to delete converted file: {output_file}")
                #     deletion_result = cleanup_file(output_file)  # Attempt to delete converted file
                #     if deletion_result:
                #         print(f"Successfully deleted converted file: {output_file}")
                #     else:
                #         print(f"Failed to delete converted file: {output_file}")

                # else:
                #     print(f"Converted file does not exist: {output_file}")

    return render_template('video_convert.html')

@app.route('/video_to_audio', methods=['GET', 'POST'])
def video_to_audio():
    """Extract audio from video files."""
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)
        
        file = request.files['file']
        
        if file and allowed_file(file.filename, ALLOWED_VIDEO_EXTENSIONS):
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            
            selected_format = request.form['format']  # Get the selected format
            output_audio_file = os.path.splitext(file_path)[0] + f'.{selected_format}'  # Use selected format
            
            codec_map = {
                'mp3': 'libmp3lame',
                'wav': 'pcm_s16le',
                'aac': 'aac',
                'flac': 'flac',
                'm4a': 'aac'  # Use AAC codec for M4A
            }
            
            # Initialize conversion logging variables
            original_file_size = os.path.getsize(file_path)
            converted_file_size = 0
            conversion_time = 0.0
            conversion_status = 'failed'  # Default status for failure
            file_id = None  # Initialize file_id
            
            try:
                # Start measuring conversion time
                start_time = time.time()
                
                # Process video to audio conversion
                video = VideoFileClip(file_path)
                audio = video.audio
                audio.write_audiofile(output_audio_file, codec=codec_map.get(selected_format, 'libmp3lame'))
                
                # End measuring conversion time
                end_time = time.time()
                conversion_time = end_time - start_time  # Calculate time in seconds

                # Get converted file size for logging
                converted_file_size = os.path.getsize(output_audio_file)

                # If everything goes well, mark conversion as successful
                conversion_status = 'successful'  

                # Log the audio file metadata and retrieve the file_id
                file_id = log_vid_to_aud_file_metadata(filename, file.content_type, selected_format, conversion_status)

                if file_id:
                    # Log the successful conversion
                    log_vid_to_aud_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status)

                # Send the converted file back to the user
                return send_file(output_audio_file, as_attachment=True)

            except Exception as e:
                # Handle conversion errors
                logging.error(f"Conversion error: {e}")
                flash(f'Error extracting audio: {e}', 'error')

                # Log the failed conversion with the existing file sizes and conversion time
                if file_id is None:
                    # Log the metadata first to get the file_id if conversion failed
                    file_id = log_vid_to_aud_file_metadata(filename, file.content_type, selected_format, conversion_status)

                log_vid_to_aud_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status)

                return redirect(request.url)
            
            finally:
                time.sleep(conv_time)
                #cleanup_file(file_path)
                if output_audio_file:
                    insert_file_for_deletion(output_audio_file)
                if file_path:
                    insert_file_for_deletion(file_path)
                # # Clean up temporary files
                # if video is not None:
                #     video.reader.close()
                #     if video.audio is not None:
                #         video.audio.reader.close_proc()

                # # Ensure files are deleted after closing video
                # cleanup_file(file_path)  # Delete original file

                # time.sleep(1)  # Wait for 10 seconds before attempting to delete the converted file

                # # Attempt to delete the converted file
                # if os.path.exists(output_audio_file):
                #     print(f"Attempting to delete converted file: {output_audio_file}")
                #     deletion_result = cleanup_file(output_audio_file)  # Attempt to delete converted file
                #     if deletion_result:
                #         print(f"Successfully deleted converted file: {output_audio_file}")
                #     else:
                #         print(f"Failed to delete converted file: {output_audio_file}")

                # else:
                #     print(f"Converted file does not exist: {output_audio_file}")
    
    return render_template('video_to_audio.html')

@app.route('/remove_audio', methods=['GET', 'POST'])
def remove_audio():
    """Remove audio from video files."""
    if request.method == 'POST':
        if 'file' not in request.files:
            flash('No file part')
            return redirect(request.url)

        file = request.files['file']

        if file and allowed_file(file.filename, ALLOWED_VIDEO_EXTENSIONS):
            # Check if the file size exceeds 100MB (100 * 1024 * 1024 bytes)
            if file.content_length > 100 * 1024 * 1024:
                flash('File size must be less than 100MB.')
                return redirect(request.url)

            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            output_file = os.path.splitext(file_path)[0] + '_no_audio.mp4'

            # Initialize conversion details
            conversion_status = 'failed'
            conversion_time = 0.0
            original_file_size = os.path.getsize(file_path)
            conversion_type = f"{file.content_type} - no_audio"  # Set your conversion type accordingly

            video = None  # Initialize video variable for cleanup

            try:
                # Start conversion and log time
                start_time = time.time()
                video = VideoFileClip(file_path)
                video_no_audio = video.without_audio()
                video_no_audio.write_videofile(
                    output_file,
                    codec='libx264',
                    bitrate='5000k',
                    preset='slow',
                    ffmpeg_params=['-crf', '18']
                )

                # Mark conversion as successful
                conversion_status = 'successful'
                conversion_time = time.time() - start_time
                converted_file_size = os.path.getsize(output_file)

                # Log conversion details to the database
                file_id = log_mute_vid_file_metadata(filename, file.content_type, 'mp4', conversion_status, conversion_type)
                log_mute_vid_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type)

                # Send the file for download
                return send_file(output_file, as_attachment=True)

            except Exception as e:
                flash(f'Error removing audio: {e}')
                # Log failed conversion
                log_mute_vid_conversion(None, original_file_size, None, conversion_time, conversion_status, conversion_type)
                return redirect(request.url)

            finally:
                time.sleep(conv_time)
                #cleanup_file(file_path)
                if output_file:
                    insert_file_for_deletion(output_file)
                if file_path:
                    insert_file_for_deletion(file_path)
                # # Clean up temporary files
                # if video is not None:
                #     video.reader.close()
                #     if video.audio is not None:
                #         video.audio.reader.close_proc()

                # # Ensure files are deleted after closing video
                # cleanup_file(file_path)  # Delete original file

                # time.sleep(1)  # Wait for 10 seconds before attempting to delete the converted file

                # # Attempt to delete the converted file
                # if os.path.exists(output_file):
                #     print(f"Attempting to delete converted file: {output_file}")
                #     deletion_result = cleanup_file(output_file)  # Attempt to delete converted file
                #     if deletion_result:
                #         print(f"Successfully deleted converted file: {output_file}")
                #     else:
                #         print(f"Failed to delete converted file: {output_file}")

                # else:
                #     print(f"Converted file does not exist: {output_file}")

    return render_template('remove_audio.html')

@app.route('/document_convert', methods=['GET', 'POST'])
def document_convert():
    """Handle document file upload and conversion."""

    def save_file(file):
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        return file_path, filename

    def handle_pdf_conversion(file_path, filename, target_format):
        try:
            if target_format == 'jpg':
                return convert_pdf_to_jpg(file_path, filename)
            elif target_format == 'txt':
                return convert_pdf_to_txt(file_path)
            elif target_format == 'docx':
                return convert_pdf_to_docx(file_path)
            elif target_format == 'excel' or target_format == 'xlsx':
                return convert_pdf_to_excel(file_path)
            elif target_format == 'csv':
                return convert_pdf_to_csv(file_path)
            else:
                flash('Unsupported PDF conversion format!', 'error')
                return redirect(request.url)
        except Exception as e:
            flash(f'Error converting PDF: {e}', 'error')
            return redirect(request.url)

    def handle_docx_conversion(file_path, filename, target_format):
        try:
            if target_format == 'pdf':
                return convert_docx_to_pdf(file_path)
            elif target_format == 'txt':
                return convert_docx_to_txt(file_path)
            elif target_format == 'excel' or target_format == 'xlsx':
                return convert_docx_to_excel(file_path)
            elif target_format == 'csv':
                return convert_docx_to_csv(file_path)
            elif target_format == 'docx':
                return send_file(file_path, as_attachment=True)
            else:
                flash('Unsupported DOCX conversion format!', 'error')
                return redirect(request.url)
        except Exception as e:
            flash(f'Error converting DOCX: {e}', 'error')
            return redirect(request.url)

    def handle_txt_conversion(file_path, target_format):
        try:
            if target_format == 'docx':
                return convert_txt_to_docx(file_path)
            elif target_format == 'pdf':
                return convert_txt_to_pdf(file_path)
            elif target_format == 'xlsx':
                return convert_txt_to_excel(file_path)
            elif target_format == 'csv':
                return convert_txt_to_csv(file_path)
            elif target_format == 'txt':
                return send_file(file_path, as_attachment=True)
            else:
                flash('Unsupported text conversion format!', 'error')
                return redirect(request.url)
        except Exception as e:
            flash(f'Error converting text file: {e}', 'error')
            return redirect(request.url)

    def handle_excel_conversion(file_path, filename, target_format):
        try:
            if target_format == 'csv':
                return convert_excel_to_csv(file_path)
            elif target_format == 'pdf':
                return convert_excel_to_pdf(file_path)
            elif target_format == 'xlsx':
                return send_file(file_path, as_attachment=True)
            elif target_format == 'docx':
                return convert_excel_to_docx(file_path)
            elif target_format == 'txt':
                return convert_excel_to_txt(file_path)
            else:
                flash('Unsupported Excel conversion format!', 'error')
                return redirect(request.url)
        except Exception as e:
            flash(f'Error converting Excel file: {e}', 'error')
            return redirect(request.url)
        
    def handle_csv_conversion(file_path, filename, target_format):
        try:
            if target_format == 'excel' or target_format == 'xlsx':
                return convert_csv_to_excel(file_path)
            elif target_format == 'pdf':
                return convert_csv_to_pdf(file_path)
            elif target_format == 'docx':
                return convert_csv_to_docx(file_path)
            elif target_format == 'txt':
                return convert_csv_to_txt(file_path)
            elif target_format == 'csv':
                return send_file(file_path, as_attachment=True)
            else:
                flash('Unsupported CSV conversion format!', 'error')
                return redirect(request.url)
        except Exception as e:
            flash(f'Error converting CSV file: {e}', 'error')
            return redirect(request.url)


    def convert_pdf_to_jpg(pdf_path, filename):
        images = convert_from_path(pdf_path)
        output_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'images')
        os.makedirs(output_dir, exist_ok=True)
        image_paths = []
        for i, image in enumerate(images):
            image_path = os.path.join(output_dir, f'{os.path.splitext(filename)[0]}_page_{i + 1}.jpg')
            image.save(image_path, 'JPEG')
            image_paths.append(image_path)
        return image_paths[0]

    def convert_pdf_to_txt(pdf_path):
        text = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        
        txt_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(pdf_path))[0]}.txt')
        with open(txt_path, 'w') as txt_file:
            txt_file.write(text)
        return txt_path

    def convert_pdf_to_docx(pdf_path):
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(pdf_path))[0]}.docx')
        cv = Converter(pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
        return docx_path

    def convert_pdf_to_excel(pdf_path):
        print(f"Starting conversion of PDF: {pdf_path}")
        
        # Define the Excel file path
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(pdf_path))[0]}.xlsx')
        print(f"Output Excel path: {excel_path}")
        
        data = []
        
        try:
            # Using pdfplumber to read the PDF
            with pdfplumber.open(pdf_path) as pdf:
                print(f"Number of pages in PDF: {len(pdf.pages)}")
                
                # Read each page of the PDF
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    print(f"Extracting text from page {page_num + 1}")
                    
                    if text:  # Only process if text is not None
                        for line in text.split('\n'):
                            stripped_line = line.strip()
                            if stripped_line:  # Check for non-empty lines
                                data.append([stripped_line])  # Add each line as a row in the Excel file
                                print(f"Added line to data: {stripped_line}")

            # Check if data is empty
            if not data:
                print("No data extracted from PDF.")
                flash('No data extracted from the PDF.', 'error')
                return redirect(request.url)

            # Create a DataFrame and write to Excel
            df = pd.DataFrame(data)
            df.to_excel(excel_path, index=False, header=False)
            print(f"Successfully written data to Excel: {excel_path}")

        except Exception as e:
            print(f"Error during PDF processing: {e}")
            flash(f'Error converting PDF: {e}', 'error')
            return redirect(request.url)

        return excel_path

    def convert_docx_to_pdf(docx_path):
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(docx_path))[0]}.pdf')
        
        try:
            print(f"Attempting to convert DOCX file: {docx_path}")
            
            if not os.path.exists(docx_path):
                print("DOCX file does not exist.")
                raise Exception(f"DOCX file not found: {docx_path}")

            print(f"Converting {docx_path} to {pdf_path}")
            docx2pdf_convert(docx_path, pdf_path)
            
            print(f"Conversion successful: {pdf_path} created.")
            return pdf_path
        
        except Exception as e:
            print(f"Error during conversion: {e}")
            raise Exception(f'Error converting DOCX to PDF: {e}')

    def convert_docx_to_txt(docx_path):
        doc = docx.Document(docx_path)
        text = '\n'.join([para.text for para in doc.paragraphs])

        txt_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(docx_path))[0]}.txt')
        with open(txt_path, 'w') as txt_file:
            txt_file.write(text)

        return txt_path

    def convert_docx_to_excel(docx_path):
        try:
            if not os.path.exists(docx_path):
                flash('The specified DOCX file does not exist!', 'error')
                return redirect(request.url)

            text = docx2txt.process(docx_path)
            lines = text.splitlines()  # Split text into lines

            # Create DataFrame
            df = pd.DataFrame(lines, columns=['Content'])
            excel_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(docx_path))[0]}.xlsx')
            df.to_excel(excel_path, index=False)
            return excel_path
        except Exception as e:
            flash(f'Error converting DOCX to Excel: {e}', 'error')
            return redirect(request.url)

    def convert_txt_to_docx(txt_path):
        doc = docx.Document()
        with open(txt_path, 'r') as txt_file:
            for line in txt_file:
                doc.add_paragraph(line.strip())

        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(txt_path))[0]}.docx')
        doc.save(docx_path)

        return docx_path

    def convert_txt_to_pdf(txt_path):
        try:
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(txt_path))[0]}.pdf')
            
            # Read content from the text file
            with open(txt_path, 'r') as file:
                content = file.read()
            
            # Create a simple HTML representation of the text content
            html_content = f"<html><body><pre>{content}</pre></body></html>"

            # Convert HTML to PDF
            pdfkit.from_string(html_content, pdf_path, configuration=pdf_config)
            return pdf_path
        except Exception as e:
            raise Exception(f'Error converting TXT to PDF: {e}')

    def convert_txt_to_excel(txt_path):
        excel_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(txt_path))[0]}.xlsx')
        data = []
        with open(txt_path, 'r') as txt_file:
            for line in txt_file:
                data.append([line.strip()])
        df = pd.DataFrame(data, columns=['Content'])
        df.to_excel(excel_path, index=False)
        return excel_path

    def convert_excel_to_csv(excel_path):
        try:
            df = pd.read_excel(excel_path)
            csv_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(excel_path))[0]}.csv')
            df.to_csv(csv_path, index=False)
            return csv_path
        except Exception as e:
            raise Exception(f'Error converting Excel to CSV: {e}')

    def convert_csv_to_excel(csv_path):
        try:
            df = pd.read_csv(csv_path)
            excel_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(csv_path))[0]}.xlsx')
            df.to_excel(excel_path, index=False)
            return excel_path
        except Exception as e:
            raise Exception(f'Error converting CSV to Excel: {e}')

    def convert_csv_to_pdf(csv_path):
        try:
            print(f"Attempting to convert CSV file: {csv_path}")
            
            if not os.path.exists(csv_path):
                print("CSV file does not exist.")
                raise Exception(f"CSV file not found: {csv_path}")
            
            print(f"Reading CSV file: {csv_path}")
            df = pd.read_csv(csv_path)
            
            print(f"DataFrame read successfully. Converting to HTML.")
            html = df.to_html(index=False)
            
            pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(csv_path))[0]}.pdf')
            
            
            # Update to pass configuration directly in the conversion call
            pdfkit.from_string(html, pdf_path, configuration=pdf_config)
            
            print(f"Conversion successful: {pdf_path} created.")
            return pdf_path
        
        except Exception as e:
            print(f"Error during conversion: {e}")
            raise Exception(f'Error converting CSV to PDF: {e}')

    def convert_csv_to_docx(csv_path):
        try:
            df = pd.read_csv(csv_path)
            doc = Document()
            table = doc.add_table(rows=1, cols=len(df.columns))
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(df.columns):
                hdr_cells[i].text = column_name
            for index, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)
            docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(csv_path))[0]}.docx')
            doc.save(docx_path)
            return docx_path
        except Exception as e:
            raise Exception(f'Error converting CSV to DOCX: {e}')

    def convert_csv_to_txt(csv_path):
        try:
            df = pd.read_csv(csv_path)
            txt_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(csv_path))[0]}.txt')
            df.to_csv(txt_path, sep='\t', index=False)  # Save as tab-separated values
            return txt_path
        except Exception as e:
            raise Exception(f'Error converting CSV to TXT: {e}')
        
    def convert_docx_to_csv(docx_path):
        # Read the DOCX file
        doc = Document(docx_path)
        data = []

        for para in doc.paragraphs:
            data.append([para.text])  # Add each paragraph as a row in the CSV

        # Define the CSV file path
        csv_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(docx_path))[0]}.csv')
        
        # Write to CSV
        with open(csv_path, mode='w', newline='', encoding='utf-8') as csv_file:
            writer = csv.writer(csv_file)
            writer.writerows(data)

        return csv_path

    def convert_txt_to_csv(txt_path):
        data = []
        
        # Read the TXT file
        with open(txt_path, 'r', encoding='utf-8') as file:
            for line in file:
                data.append([line.strip()])  # Add each line as a row in the CSV

        # Define the CSV file path
        csv_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(txt_path))[0]}.csv')
        
        # Write to CSV
        with open(csv_path, mode='w', newline='', encoding='utf-8') as csv_file:
            writer = csv.writer(csv_file)
            writer.writerows(data)

        return csv_path

    def convert_pdf_to_csv(pdf_path):
        reader = PdfReader(pdf_path)
        data = []

        # Read each page of the PDF
        for page in reader.pages:
            text = page.extract_text()
            if text:
                for line in text.split('\n'):
                    data.append([line.strip()])  # Add each line as a row in the CSV

        # Define the CSV file path
        csv_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(pdf_path))[0]}.csv')
        
        # Write to CSV
        with open(csv_path, mode='w', newline='', encoding='utf-8') as csv_file:
            writer = csv.writer(csv_file)
            writer.writerows(data)

        return csv_path

    def convert_excel_to_pdf(excel_path):
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(excel_path))[0]}.pdf')

        if not os.path.exists(excel_path):
            raise Exception(f"Excel file not found: {excel_path}")

        try:
            df = pd.read_excel(excel_path)
            html = df.to_html(index=False)

            # Use the global pdf_config for conversion
            pdfkit.from_string(html, pdf_path, configuration=pdf_config)
            return pdf_path
        except Exception as e:
            print(f"PDF Conversion Error: {e}")  # Debugging statement
            raise Exception(f'Error converting Excel to PDF: {e}')

    def convert_excel_to_docx(excel_path):
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(excel_path))[0]}.docx')
        df = pd.read_excel(excel_path)
        doc = docx.Document()
        for _, row in df.iterrows():
            doc.add_paragraph(','.join(str(cell) for cell in row))
        doc.save(docx_path)
        return docx_path

    def convert_excel_to_txt(excel_path):
        txt_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{os.path.splitext(os.path.basename(excel_path))[0]}.txt')
        df = pd.read_excel(excel_path)
        with open(txt_path, 'w') as txt_file:
            for _, row in df.iterrows():
                txt_file.write(','.join(str(cell) for cell in row) + '\n')
        return txt_path

    if request.method == 'POST':
        file = request.files.get('file')
        if file and allowed_file(file.filename, ALLOWED_DOCUMENT_EXTENSIONS):
            # Get the secure filename and save the file
            filename = secure_filename(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)

            # Extract the file extension from the filename
            original_format = os.path.splitext(filename)[1].replace('.', '')
            target_format = request.form.get('format')
            output_document_file = os.path.splitext(file_path)[0] + f'.{target_format}'

            # Initialize variables for conversion logging
            original_file_size = os.path.getsize(file_path)
            converted_file_size = None
            conversion_time = None
            conversion_status = 'failed'  # Default status for failure
            conversion_type = f"{original_format} to {target_format}"  # Set your conversion type accordingly

            try:
                # Start measuring conversion time
                start_time = time.time()

                # Process document conversion (replace with your document conversion logic)
                if filename.lower().endswith('.pdf'):
                    print("Handling PDF conversion...")
                    output_document_file = handle_pdf_conversion(file_path, filename, target_format)
                elif filename.lower().endswith('.docx'):
                    print("Handling DOCX conversion...")
                    output_document_file = handle_docx_conversion(file_path, filename, target_format)
                elif filename.lower().endswith('.txt'):
                    print("Handling TXT conversion...")
                    output_document_file = handle_txt_conversion(file_path, target_format)
                elif filename.lower().endswith('.xlsx'):
                    print("Handling Excel conversion...")
                    output_document_file = handle_excel_conversion(file_path, filename, target_format)
                elif filename.lower().endswith('.csv'):
                    print("Handling CSV conversion...")
                    output_document_file = handle_csv_conversion(file_path, filename, target_format)
                else:
                    flash('Unsupported file type!', 'error')
                    print(f'Unsupported file type: {filename}')
                    return redirect(request.url)

                # End measuring conversion time
                end_time = time.time()
                conversion_time = end_time - start_time  # Calculate time in seconds as float

                # Get file sizes for logging
                converted_file_size = os.path.getsize(output_document_file)

                # If everything goes well, mark conversion as successful
                conversion_status = 'successful'

                # Log the file metadata
                log_doc_file_metadata(filename, original_format, target_format, conversion_type)

                # Log the successful conversion
                log_doc_conversion(conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status)

                # Send the converted file back to the user
                return send_file(output_document_file, as_attachment=True, download_name=f'{filename}.{target_format}')

            except Exception as e:
                # Handle conversion errors
                logging.error(f"Conversion error: {e}")
                flash(f'Error processing document file: {e}', 'error')

                # Log the failed conversion (no need to log metadata since it wasn't successful)
                log_doc_conversion(conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status)

                return redirect(request.url)
            
            finally:
                # Ensure files are deleted after conversion
                time.sleep(conv_time)
                if output_document_file and isinstance(output_document_file, str):  # Make sure result is a file path
                    insert_file_for_deletion(output_document_file)  # Insert the file path to delete later
                if file_path:
                    insert_file_for_deletion(file_path)

        else:
            flash('Invalid file format. Please upload a valid document file.', 'error')

    return render_template('document_convert.html')

def create_connection():
    """Create a database connection to the MySQL database."""
    connection = None
    try:
        connection = mysql.connector.connect(
            host='localhost',
            user='tiesetso',  # Your MySQL username
            password='Mat67087@testing',  # Your MySQL password
            database='utility_tools'  # Your MySQL database name
        )
        logging.info("Connection to MySQL DB successful")
    except Error as e:
        logging.error(f"The error '{e}' occurred")
    return connection

@app.route('/feedback', methods=['GET', 'POST'])
def feedback():
    if request.method == 'POST':
        message = request.form['message']

        # Insert feedback into the database
        db = create_connection()  # Create a new connection
        cursor = db.cursor()
        cursor.execute("INSERT INTO feedback (message) VALUES (%s)", (message,))
        db.commit()
        cursor.close()
        db.close()  # Close the database connection

        flash('Thank you for your feedback!', 'success')
        return redirect('/feedback')

    return render_template('feedback.html')

@app.route('/submit_feedback', methods=['POST'])
def submit_feedback():
    message = request.form['message']
    db = create_connection()  # Create a new connection
    if db is None:
        flash('Database connection failed!', 'error')
        return redirect(url_for('feedback'))  # Redirect to the feedback page

    try:
        # Insert message into the database
        cursor = db.cursor()
        cursor.execute("INSERT INTO feedback (message) VALUES (%s)", (message,))
        db.commit()
        cursor.close()
        flash('Thank you! We have successfully received your message and hope to improve based on your feedback.', 'success')
    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'error')
    finally:
        db.close()  # Close the connection after use

    return redirect(url_for('feedback'))  # Redirect to the feedback page

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        message = request.form['message']

        # Insert contact form data into the database
        db = create_connection()  # Create a new connection
        cursor = db.cursor()
        cursor.execute("INSERT INTO contacts (name, email, message) VALUES (%s, %s, %s)", (name, email, message))
        db.commit()
        cursor.close()
        db.close()  # Close the database connection

        flash('Thank you, {}! Your message has been received successfully. We will get back to you soon.'.format(name), 'success')
        return redirect('/contact')

    return render_template('contact.html')

@app.route('/terms-and-conditions', methods=['GET'])
def terms_and_conditions():
    return render_template('terms_and_conditions.html')

@app.route('/about_us', methods=['GET'])
def about_us():
    return render_template('about_us.html')

@app.route('/conversion-progress')
def conversion_progress():
    """Return the conversion progress."""
    return jsonify(progress=session.get('progress', 0))

@app.route('/download-progress')
def download_progress():
    """Return the download progress."""
    return jsonify(progress=session.get('download_progress', 0))

### MySQL code for the Admin Dashboard###

def log_doc_conversion(conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status):
    """Log conversion type to the doc_conversions table."""
    conn = create_connection()
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO doc_conversions (conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status)
            VALUES (%s, %s, %s, %s, %s)
        """
        cursor.execute(query, (conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status))
        conn.commit()
        cursor.close()
        conn.close()

def log_doc_file_metadata(file_name, original_format, target_format, conversion_type):
    """Log file metadata to the doc_file_metadata table."""
    conn = create_connection()
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO doc_file_metadata (file_name, original_format, target_format, conversion_type)
            VALUES (%s, %s, %s, %s)
        """
        cursor.execute(query, (file_name, original_format, target_format, conversion_type))
        conn.commit()
        cursor.close()
        conn.close()

def log_audio_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type):
    """Log audio conversion details to the aud_conversions table."""
    conn = create_connection()
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO aud_conversions (file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type)
            VALUES (%s, %s, %s, %s, %s, %s)
        """
        cursor.execute(query, (file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type))
        conn.commit()
        cursor.close()
        conn.close()

def log_audio_file_metadata(file_name, original_format, target_format, conversion_status, conversion_type):
    """Log file metadata to the aud_file_metadata table and return the file_id."""
    conn = create_connection()
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO aud_file_metadata (file_name, original_format, target_format, conversion_status, conversion_type)
            VALUES (%s, %s, %s, %s, %s)
        """
        cursor.execute(query, (file_name, original_format, target_format, conversion_status, conversion_type))
        conn.commit()
        file_id = cursor.lastrowid  # Get the last inserted ID
        cursor.close()
        conn.close()
        return file_id  # Return the file_id
    return None  # Return None if the connection fails

def log_vid_to_aud_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status):
    """Log video to audio conversion details to the vid_to_aud_conversions table."""
    conn = create_connection()  # Ensure this function establishes a connection to your database
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO vid_to_aud_conversions (file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type)
            VALUES (%s, %s, %s, %s, %s, 'video_to_audio')
        """
        cursor.execute(query, (file_id, original_file_size, converted_file_size, conversion_time, conversion_status))
        conn.commit()
        cursor.close()
        conn.close()

def log_vid_to_aud_file_metadata(file_name, original_format, target_format, conversion_status):
    """Log video file metadata to the vid_to_aud_file_metadata table and return the file_id."""
    conn = create_connection()  # Ensure this function establishes a connection to your database
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO vid_to_aud_file_metadata (file_name, original_format, target_format, conversion_status, conversion_type)
            VALUES (%s, %s, %s, %s, 'video_to_audio')
        """
        cursor.execute(query, (file_name, original_format, target_format, conversion_status))
        conn.commit()
        file_id = cursor.lastrowid  # Get the last inserted ID
        cursor.close()
        conn.close()
        return file_id  # Return the file_id
    return None  # Return None if the connection fails

def log_vid_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type):
    """Log video conversion details to the vid_conversions table."""
    conn = create_connection()  # Ensure this function establishes a connection to your database
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO vid_conversions (file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type)
            VALUES (%s, %s, %s, %s, %s, %s)
        """
        cursor.execute(query, (file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type))
        conn.commit()
        cursor.close()
        conn.close()

def log_vid_file_metadata(file_name, original_format, target_format, conversion_status, conversion_type):
    """Log video file metadata to the vid_file_metadata table and return the file_id."""
    conn = create_connection()  # Ensure this function establishes a connection to your database
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO vid_file_metadata (file_name, original_format, target_format, conversion_status, conversion_type)
            VALUES (%s, %s, %s, %s, %s)
        """
        cursor.execute(query, (file_name, original_format, target_format, conversion_status, conversion_type))
        conn.commit()
        file_id = cursor.lastrowid  # Get the last inserted ID
        cursor.close()
        conn.close()
        return file_id  # Return the file_id
    return None  # Return None if the connection fails

def log_mute_vid_conversion(file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type):
    """Log muted video conversion details to the mute_vid_conversions table."""
    conn = create_connection()  # Ensure this function establishes a connection to your database
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO mute_vid_conversions (file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type)
            VALUES (%s, %s, %s, %s, %s, %s)
        """
        cursor.execute(query, (file_id, original_file_size, converted_file_size, conversion_time, conversion_status, conversion_type))
        conn.commit()
        cursor.close()
        conn.close()

def log_mute_vid_file_metadata(file_name, original_format, target_format, conversion_status, conversion_type):
    """Log muted video file metadata to the mute_vid_file_metadata table and return the file_id."""
    conn = create_connection()  # Ensure this function establishes a connection to your database
    if conn:
        cursor = conn.cursor()
        query = """
            INSERT INTO mute_vid_file_metadata (file_name, original_format, target_format, conversion_status, conversion_type)
            VALUES (%s, %s, %s, %s, %s)
        """
        cursor.execute(query, (file_name, original_format, target_format, conversion_status, conversion_type))
        conn.commit()
        file_id = cursor.lastrowid  # Get the last inserted ID
        cursor.close()
        conn.close()
        return file_id  # Return the file_id
    return None  # Return None if the connection fails

def log_img_conversion(conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status, file_id):
    """Log conversion details to the img_conversions table.

    Args:
        conversion_type (str): The type of conversion (e.g., 'JPEG to PNG').
        original_file_size (int): The size of the original file in bytes.
        converted_file_size (int): The size of the converted file in bytes.
        conversion_time (float): The time taken for conversion in seconds.
        conversion_status (str): The status of the conversion (e.g., 'successful', 'failed').
    """
    try:
        with create_connection() as conn:
            with conn.cursor() as cursor:
                query = """
                    INSERT INTO img_conversions (conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status, file_id)
                    VALUES (%s, %s, %s, %s, %s, %s)
                """
                cursor.execute(query, (conversion_type, original_file_size, converted_file_size, conversion_time, conversion_status, file_id))
                conn.commit()
    except Exception as e:
        print(f'Error logging image conversion: {e}')

def log_img_file_metadata(file_name, original_format, target_format, conversion_type, conversion_status):
    """Log file metadata to the img_file_metadata table.

    Args:
        file_name (str): The name of the file being converted.
        original_format (str): The format of the original file (e.g., 'JPEG').
        target_format (str): The format of the converted file (e.g., 'PNG').
        conversion_type (str): The type of conversion (e.g., 'resize', 'format change').
        conversion_status (str): The status of the conversion (e.g., 'success', 'failure').

    Returns:
        int: The ID of the newly inserted record.
    """
    try:
        with create_connection() as conn:
            with conn.cursor() as cursor:
                query = """
                    INSERT INTO img_file_metadata (file_name, original_format, target_format, conversion_type, conversion_status)
                    VALUES (%s, %s, %s, %s, %s)
                """
                cursor.execute(query, (file_name, original_format, target_format, conversion_type, conversion_status))
                conn.commit()
                return cursor.lastrowid  # Return the ID of the newly inserted record
    except Exception as e:
        print(f'Error logging file metadata: {e}')
        return None  # Return None if there's an error

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        connection = create_connection()
        cursor = connection.cursor()
        cursor.execute("SELECT password FROM users WHERE username = %s", (username,))
        user = cursor.fetchone()
        
        if user and bcrypt.checkpw(password.encode('utf-8'), user[0].encode('utf-8')):
            session['username'] = username
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password', 'danger')

    return render_template('login.html')


def get_conversion_data(period):
    conn = create_connection()
    cursor = conn.cursor(dictionary=True)

    # Calculate date range based on the selected period
    if period == "today":
        start_date = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = datetime.now().replace(hour=23, minute=59, second=59, microsecond=999999)
    elif period == "week":
        start_date = datetime.now() - timedelta(days=datetime.now().weekday())
        end_date = datetime.now() + timedelta(days=(6 - datetime.now().weekday()))
    elif period == "month":
        start_date = (datetime.now().replace(day=1))
        end_date = (datetime.now().replace(day=1) + timedelta(days=32)).replace(day=1) - timedelta(seconds=1)

    # Combine conversion data from all relevant tables
    queries = [
        f"""
        SELECT 'audio' AS file_type, 
               SUM(CASE WHEN conversion_status = 'successful' THEN 1 ELSE 0 END) AS successful,
               SUM(CASE WHEN conversion_status = 'failed' THEN 1 ELSE 0 END) AS failed,
               MIN(conversion_time) AS fastest_time,
               MAX(conversion_time) AS slowest_time
        FROM aud_conversions
        WHERE conversion_date BETWEEN %s AND %s
        """,
        f"""
        SELECT 'video' AS file_type, 
               SUM(CASE WHEN conversion_status = 'successful' THEN 1 ELSE 0 END) AS successful,
               SUM(CASE WHEN conversion_status = 'failed' THEN 1 ELSE 0 END) AS failed,
               MIN(conversion_time) AS fastest_time,
               MAX(conversion_time) AS slowest_time
        FROM vid_conversions
        WHERE conversion_date BETWEEN %s AND %s
        """,
        f"""
        SELECT 'document' AS file_type, 
               SUM(CASE WHEN conversion_status = 'successful' THEN 1 ELSE 0 END) AS successful,
               SUM(CASE WHEN conversion_status = 'failed' THEN 1 ELSE 0 END) AS failed,
               MIN(conversion_time) AS fastest_time,
               MAX(conversion_time) AS slowest_time
        FROM doc_conversions
        WHERE conversion_date BETWEEN %s AND %s
        """,
        f"""
        SELECT 'image' AS file_type, 
               SUM(CASE WHEN conversion_status = 'successful' THEN 1 ELSE 0 END) AS successful,
               SUM(CASE WHEN conversion_status = 'failed' THEN 1 ELSE 0 END) AS failed,
               MIN(conversion_time) AS fastest_time,
               MAX(conversion_time) AS slowest_time
        FROM img_conversions
        WHERE conversion_date BETWEEN %s AND %s
        """,
        f"""
        SELECT 'mute_video' AS file_type, 
               SUM(CASE WHEN conversion_status = 'successful' THEN 1 ELSE 0 END) AS successful,
               SUM(CASE WHEN conversion_status = 'failed' THEN 1 ELSE 0 END) AS failed,
               MIN(conversion_time) AS fastest_time,
               MAX(conversion_time) AS slowest_time
        FROM mute_vid_conversions
        WHERE conversion_date BETWEEN %s AND %s
        """,
        f"""
        SELECT 'vid_to_audio' AS file_type, 
               SUM(CASE WHEN conversion_status = 'successful' THEN 1 ELSE 0 END) AS successful,
               SUM(CASE WHEN conversion_status = 'failed' THEN 1 ELSE 0 END) AS failed,
               MIN(conversion_time) AS fastest_time,
               MAX(conversion_time) AS slowest_time
        FROM vid_to_aud_conversions
        WHERE conversion_date BETWEEN %s AND %s
        """
    ]

    data = []
    for query in queries:
        cursor.execute(query, (start_date, end_date))
        results = cursor.fetchall()
        for row in results:
            data.append({
                'fileType': row['file_type'],
                'successful': row['successful'],
                'failed': row['failed'],
                'fastestTime': row['fastest_time'],
                'slowestTime': row['slowest_time']
            })

    cursor.close()
    conn.close()

    return data

@app.route('/dashboard')
def dashboard():
    period = request.args.get('period', 'today')
    conversion_data = get_conversion_data(period)
    print(conversion_data)
    # Check if the request is coming from an AJAX call
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        # Return conversion data as JSON for the AJAX request
        return jsonify(conversion_data)
    
    # For regular page loads, render the HTML template
    return render_template('dashboard.html', conversion_data=conversion_data, period=period)

@app.route('/documents')
def documents():
    if 'username' not in session:
        return redirect(url_for('login'))

    connection = create_connection()
    cursor = connection.cursor(named_tuple=True)

    # Fetch data from the 'doc_file_metadata' table
    cursor.execute("SELECT * FROM doc_file_metadata")
    document_data = cursor.fetchall()

    # Fetch conversion statistics
    cursor.execute("SELECT COUNT(*) AS total_conversion_attempts FROM doc_conversions")
    total_attempts = cursor.fetchone().total_conversion_attempts

    cursor.execute("SELECT COUNT(*) AS total_successful_conversions FROM doc_conversions WHERE conversion_status = 'successful'")
    total_successful = cursor.fetchone().total_successful_conversions

    cursor.execute("SELECT COUNT(*) AS total_failed_conversions FROM doc_conversions WHERE conversion_status = 'failed'")
    total_failed = cursor.fetchone().total_failed_conversions

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count DESC
        LIMIT 1
    """)
    most_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count ASC
        LIMIT 1
    """)
    least_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT all_types.conversion_type
        FROM (SELECT DISTINCT conv.conversion_type FROM doc_conversions conv) AS all_types
        LEFT JOIN (SELECT DISTINCT conv.conversion_type FROM doc_conversions conv WHERE conv.conversion_status = 'successful') AS successful_types
        ON all_types.conversion_type = successful_types.conversion_type
        WHERE successful_types.conversion_type IS NULL
    """)
    never_used_conversion = cursor.fetchall()

    cursor.execute(""" 
        SELECT conv.conversion_type, MAX(conv.conversion_time) AS max_conversion_time
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY max_conversion_time DESC
        LIMIT 1
    """)
    slowest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, MIN(conv.conversion_time) AS min_conversion_time
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY min_conversion_time ASC
        LIMIT 1
    """)
    fastest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 * 60 + HOUR(NOW()) * 60 + MINUTE(NOW())) AS conversions_per_minute
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_minute = cursor.fetchone().conversions_per_minute

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 + HOUR(NOW())) AS conversions_per_hour
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_hour = cursor.fetchone().conversions_per_hour

    cursor.execute(""" 
        SELECT COUNT(*) / DATEDIFF(NOW(), MIN(conv.conversion_date)) AS conversions_per_day
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_day = cursor.fetchone().conversions_per_day

    # Close the cursor and connection
    cursor.close()
    connection.close()

    # Render the documents template with the data from all tables and conversion statistics
    return render_template('documents.html', 
                           document_data=document_data,
                           total_attempts=total_attempts,
                           total_successful=total_successful,
                           total_failed=total_failed,
                           most_frequent_conversion=most_frequent_conversion,
                           least_frequent_conversion=least_frequent_conversion,
                           never_used_conversion=never_used_conversion,
                           slowest_conversion=slowest_conversion,
                           fastest_conversion=fastest_conversion,
                           conversions_per_minute=conversions_per_minute,
                           conversions_per_hour=conversions_per_hour,
                           conversions_per_day=conversions_per_day)

@app.route('/audio')
def audio():
    if 'username' not in session:
        return redirect(url_for('login'))

    connection = create_connection()
    cursor = connection.cursor(named_tuple=True)

    # Fetch data from the 'aud_file_metadata' table
    cursor.execute("SELECT * FROM aud_file_metadata")
    audio_data = cursor.fetchall()

    # Fetch conversion statistics from 'aud_conversions'
    cursor.execute("SELECT COUNT(*) AS total_conversion_attempts FROM aud_conversions")
    total_attempts = cursor.fetchone().total_conversion_attempts

    cursor.execute("SELECT COUNT(*) AS total_successful_conversions FROM aud_conversions WHERE conversion_status = 'successful'")
    total_successful = cursor.fetchone().total_successful_conversions

    cursor.execute("SELECT COUNT(*) AS total_failed_conversions FROM aud_conversions WHERE conversion_status = 'failed'")
    total_failed = cursor.fetchone().total_failed_conversions

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM aud_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count DESC
        LIMIT 1
    """)
    most_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM aud_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count ASC
        LIMIT 1
    """)
    least_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT all_types.conversion_type
        FROM (SELECT DISTINCT conv.conversion_type FROM aud_conversions conv) AS all_types
        LEFT JOIN (SELECT DISTINCT conv.conversion_type FROM aud_conversions conv WHERE conv.conversion_status = 'successful') AS successful_types
        ON all_types.conversion_type = successful_types.conversion_type
        WHERE successful_types.conversion_type IS NULL
    """)
    never_used_conversion = cursor.fetchall()

    cursor.execute(""" 
        SELECT conv.conversion_type, MAX(conv.conversion_time) AS max_conversion_time
        FROM aud_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY max_conversion_time DESC
        LIMIT 1
    """)
    slowest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, MIN(conv.conversion_time) AS min_conversion_time
        FROM aud_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY min_conversion_time ASC
        LIMIT 1
    """)
    fastest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 * 60 + HOUR(NOW()) * 60 + MINUTE(NOW())) AS conversions_per_minute
        FROM aud_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_minute = cursor.fetchone().conversions_per_minute

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 + HOUR(NOW())) AS conversions_per_hour
        FROM aud_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_hour = cursor.fetchone().conversions_per_hour

    cursor.execute(""" 
        SELECT COUNT(*) / DATEDIFF(NOW(), MIN(conv.conversion_date)) AS conversions_per_day
        FROM aud_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_day = cursor.fetchone().conversions_per_day

    # Close the cursor and connection
    cursor.close()
    connection.close()

    # Render the audio.html template with the data from all tables and conversion statistics
    return render_template('audio.html', 
                           audio_data=audio_data,
                           total_attempts=total_attempts,
                           total_successful=total_successful,
                           total_failed=total_failed,
                           most_frequent_conversion=most_frequent_conversion,
                           least_frequent_conversion=least_frequent_conversion,
                           never_used_conversion=never_used_conversion,
                           slowest_conversion=slowest_conversion,
                           fastest_conversion=fastest_conversion,
                           conversions_per_minute=conversions_per_minute,
                           conversions_per_hour=conversions_per_hour,
                           conversions_per_day=conversions_per_day)

@app.route('/vid_to_audio')
def vid_to_audio():
    if 'username' not in session:
        return redirect(url_for('login'))

    connection = create_connection()
    cursor = connection.cursor(named_tuple=True)

    # Fetch data from the 'vid_to_aud_file_metadata' table
    cursor.execute("SELECT * FROM vid_to_aud_file_metadata")
    video_data = cursor.fetchall()

    # Fetch conversion statistics from 'vid_to_aud_conversions'
    cursor.execute("SELECT COUNT(*) AS total_conversion_attempts FROM vid_to_aud_conversions")
    total_attempts = cursor.fetchone().total_conversion_attempts

    cursor.execute("SELECT COUNT(*) AS total_successful_conversions FROM vid_to_aud_conversions WHERE conversion_status = 'successful'")
    total_successful = cursor.fetchone().total_successful_conversions

    cursor.execute("SELECT COUNT(*) AS total_failed_conversions FROM vid_to_aud_conversions WHERE conversion_status = 'failed'")
    total_failed = cursor.fetchone().total_failed_conversions

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM vid_to_aud_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count DESC
        LIMIT 1
    """)
    most_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM vid_to_aud_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count ASC
        LIMIT 1
    """)
    least_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT all_types.conversion_type
        FROM (SELECT DISTINCT conv.conversion_type FROM vid_to_aud_conversions conv) AS all_types
        LEFT JOIN (SELECT DISTINCT conv.conversion_type FROM vid_to_aud_conversions conv WHERE conv.conversion_status = 'successful') AS successful_types
        ON all_types.conversion_type = successful_types.conversion_type
        WHERE successful_types.conversion_type IS NULL
    """)
    never_used_conversion = cursor.fetchall()

    cursor.execute(""" 
        SELECT conv.conversion_type, MAX(conv.conversion_time) AS max_conversion_time
        FROM vid_to_aud_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY max_conversion_time DESC
        LIMIT 1
    """)
    slowest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, MIN(conv.conversion_time) AS min_conversion_time
        FROM vid_to_aud_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY min_conversion_time ASC
        LIMIT 1
    """)
    fastest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 * 60 + HOUR(NOW()) * 60 + MINUTE(NOW())) AS conversions_per_minute
        FROM vid_to_aud_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_minute = cursor.fetchone().conversions_per_minute

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 + HOUR(NOW())) AS conversions_per_hour
        FROM vid_to_aud_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_hour = cursor.fetchone().conversions_per_hour

    cursor.execute(""" 
        SELECT COUNT(*) / DATEDIFF(NOW(), MIN(conv.conversion_date)) AS conversions_per_day
        FROM vid_to_aud_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_day = cursor.fetchone().conversions_per_day

    # Close the cursor and connection
    cursor.close()
    connection.close()

    # Render the vid_to_audio.html template with the data from all tables and conversion statistics
    return render_template('vid_to_audio.html', 
                           video_data=video_data,
                           total_attempts=total_attempts,
                           total_successful=total_successful,
                           total_failed=total_failed,
                           most_frequent_conversion=most_frequent_conversion,
                           least_frequent_conversion=least_frequent_conversion,
                           never_used_conversion=never_used_conversion,
                           slowest_conversion=slowest_conversion,
                           fastest_conversion=fastest_conversion,
                           conversions_per_minute=conversions_per_minute,
                           conversions_per_hour=conversions_per_hour,
                           conversions_per_day=conversions_per_day)

@app.route('/video')
def video():
    if 'username' not in session:
        return redirect(url_for('login'))

    connection = create_connection()
    cursor = connection.cursor(named_tuple=True)

    # Fetch data from the 'vid_file_metadata' table
    cursor.execute("SELECT * FROM vid_file_metadata")
    video_data = cursor.fetchall()

    # Fetch conversion statistics from 'vid_conversions'
    cursor.execute("SELECT COUNT(*) AS total_conversion_attempts FROM vid_conversions")
    total_attempts = cursor.fetchone().total_conversion_attempts

    cursor.execute("SELECT COUNT(*) AS total_successful_conversions FROM vid_conversions WHERE conversion_status = 'successful'")
    total_successful = cursor.fetchone().total_successful_conversions

    cursor.execute("SELECT COUNT(*) AS total_failed_conversions FROM vid_conversions WHERE conversion_status = 'failed'")
    total_failed = cursor.fetchone().total_failed_conversions

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM vid_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count DESC
        LIMIT 1
    """)
    most_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM vid_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count ASC
        LIMIT 1
    """)
    least_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT all_types.conversion_type
        FROM (SELECT DISTINCT conv.conversion_type FROM vid_conversions conv) AS all_types
        LEFT JOIN (SELECT DISTINCT conv.conversion_type FROM vid_conversions conv WHERE conv.conversion_status = 'successful') AS successful_types
        ON all_types.conversion_type = successful_types.conversion_type
        WHERE successful_types.conversion_type IS NULL
    """)
    never_used_conversion = cursor.fetchall()

    cursor.execute(""" 
        SELECT conv.conversion_type, MAX(conv.conversion_time) AS max_conversion_time
        FROM vid_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY max_conversion_time DESC
        LIMIT 1
    """)
    slowest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, MIN(conv.conversion_time) AS min_conversion_time
        FROM vid_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY min_conversion_time ASC
        LIMIT 1
    """)
    fastest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 * 60 + HOUR(NOW()) * 60 + MINUTE(NOW())) AS conversions_per_minute
        FROM vid_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_minute = cursor.fetchone().conversions_per_minute

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 + HOUR(NOW())) AS conversions_per_hour
        FROM vid_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_hour = cursor.fetchone().conversions_per_hour

    cursor.execute(""" 
        SELECT COUNT(*) / DATEDIFF(NOW(), MIN(conv.conversion_date)) AS conversions_per_day
        FROM vid_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_day = cursor.fetchone().conversions_per_day

    # Close the cursor and connection
    cursor.close()
    connection.close()

    # Render the video.html template with the data from all tables and conversion statistics
    return render_template('video.html', 
                           video_data=video_data,
                           total_attempts=total_attempts,
                           total_successful=total_successful,
                           total_failed=total_failed,
                           most_frequent_conversion=most_frequent_conversion,
                           least_frequent_conversion=least_frequent_conversion,
                           never_used_conversion=never_used_conversion,
                           slowest_conversion=slowest_conversion,
                           fastest_conversion=fastest_conversion,
                           conversions_per_minute=conversions_per_minute,
                           conversions_per_hour=conversions_per_hour,
                           conversions_per_day=conversions_per_day)

@app.route('/mute_video')
def mute_video():
    if 'username' not in session:
        return redirect(url_for('login'))

    connection = create_connection()
    cursor = connection.cursor(named_tuple=True)

    # Fetch data from the 'mute_vid_file_metadata' table
    cursor.execute("SELECT * FROM mute_vid_file_metadata")
    video_data = cursor.fetchall()

    # Fetch conversion statistics from 'mute_vid_conversions'
    cursor.execute("SELECT COUNT(*) AS total_conversion_attempts FROM mute_vid_conversions")
    total_attempts = cursor.fetchone().total_conversion_attempts

    cursor.execute("SELECT COUNT(*) AS total_successful_conversions FROM mute_vid_conversions WHERE conversion_status = 'successful'")
    total_successful = cursor.fetchone().total_successful_conversions

    cursor.execute("SELECT COUNT(*) AS total_failed_conversions FROM mute_vid_conversions WHERE conversion_status = 'failed'")
    total_failed = cursor.fetchone().total_failed_conversions

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM mute_vid_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count DESC
        LIMIT 1
    """)
    most_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM mute_vid_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count ASC
        LIMIT 1
    """)
    least_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT all_types.conversion_type
        FROM (SELECT DISTINCT conv.conversion_type FROM mute_vid_conversions conv) AS all_types
        LEFT JOIN (SELECT DISTINCT conv.conversion_type FROM mute_vid_conversions conv WHERE conv.conversion_status = 'successful') AS successful_types
        ON all_types.conversion_type = successful_types.conversion_type
        WHERE successful_types.conversion_type IS NULL
    """)
    never_used_conversion = cursor.fetchall()

    cursor.execute(""" 
        SELECT conv.conversion_type, MAX(conv.conversion_time) AS max_conversion_time
        FROM mute_vid_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY max_conversion_time DESC
        LIMIT 1
    """)
    slowest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, MIN(conv.conversion_time) AS min_conversion_time
        FROM mute_vid_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY min_conversion_time ASC
        LIMIT 1
    """)
    fastest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 * 60 + HOUR(NOW()) * 60 + MINUTE(NOW())) AS conversions_per_minute
        FROM mute_vid_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_minute = cursor.fetchone().conversions_per_minute

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 + HOUR(NOW())) AS conversions_per_hour
        FROM mute_vid_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_hour = cursor.fetchone().conversions_per_hour

    cursor.execute(""" 
        SELECT COUNT(*) / DATEDIFF(NOW(), MIN(conv.conversion_date)) AS conversions_per_day
        FROM mute_vid_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_day = cursor.fetchone().conversions_per_day

    # Close the cursor and connection
    cursor.close()
    connection.close()

    # Render the mute_video.html template with the data from all tables and conversion statistics
    return render_template('mute_video.html', 
                           video_data=video_data,
                           total_attempts=total_attempts,
                           total_successful=total_successful,
                           total_failed=total_failed,
                           most_frequent_conversion=most_frequent_conversion,
                           least_frequent_conversion=least_frequent_conversion,
                           never_used_conversion=never_used_conversion,
                           slowest_conversion=slowest_conversion,
                           fastest_conversion=fastest_conversion,
                           conversions_per_minute=conversions_per_minute,
                           conversions_per_hour=conversions_per_hour,
                           conversions_per_day=conversions_per_day)

@app.route('/image')
def image():
    if 'username' not in session:
        return redirect(url_for('login'))

    connection = create_connection()
    cursor = connection.cursor(named_tuple=True)

    # Fetch data from the 'img_file_metadata' table
    cursor.execute("SELECT * FROM img_file_metadata")
    image_data = cursor.fetchall()

    # Fetch conversion statistics from 'img_conversions'
    cursor.execute("SELECT COUNT(*) AS total_conversion_attempts FROM img_conversions")
    total_attempts = cursor.fetchone().total_conversion_attempts

    cursor.execute("SELECT COUNT(*) AS total_successful_conversions FROM img_conversions WHERE conversion_status = 'successful'")
    total_successful = cursor.fetchone().total_successful_conversions

    cursor.execute("SELECT COUNT(*) AS total_failed_conversions FROM img_conversions WHERE conversion_status = 'failed'")
    total_failed = cursor.fetchone().total_failed_conversions

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM img_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count DESC
        LIMIT 1
    """)
    most_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM img_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY count ASC
        LIMIT 1
    """)
    least_frequent_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT all_types.conversion_type
        FROM (SELECT DISTINCT conv.conversion_type FROM img_conversions conv) AS all_types
        LEFT JOIN (SELECT DISTINCT conv.conversion_type FROM img_conversions conv WHERE conv.conversion_status = 'successful') AS successful_types
        ON all_types.conversion_type = successful_types.conversion_type
        WHERE successful_types.conversion_type IS NULL
    """)
    never_used_conversion = cursor.fetchall()

    cursor.execute(""" 
        SELECT conv.conversion_type, MAX(conv.conversion_time) AS max_conversion_time
        FROM img_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY max_conversion_time DESC
        LIMIT 1
    """)
    slowest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT conv.conversion_type, MIN(conv.conversion_time) AS min_conversion_time
        FROM img_conversions conv
        WHERE conv.conversion_status = 'successful'
        GROUP BY conv.conversion_type
        ORDER BY min_conversion_time ASC
        LIMIT 1
    """)
    fastest_conversion = cursor.fetchone()

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 * 60 + HOUR(NOW()) * 60 + MINUTE(NOW())) AS conversions_per_minute
        FROM img_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_minute = cursor.fetchone().conversions_per_minute

    cursor.execute(""" 
        SELECT COUNT(*) / (DATEDIFF(NOW(), MIN(conv.conversion_date)) * 24 + HOUR(NOW())) AS conversions_per_hour
        FROM img_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_hour = cursor.fetchone().conversions_per_hour

    cursor.execute(""" 
        SELECT COUNT(*) / DATEDIFF(NOW(), MIN(conv.conversion_date)) AS conversions_per_day
        FROM img_conversions conv
        WHERE conv.conversion_status = 'successful'
    """)
    conversions_per_day = cursor.fetchone().conversions_per_day

    # Close the cursor and connection
    cursor.close()
    connection.close()

    # Render the mute_image.html template with the data from all tables and conversion statistics
    return render_template('image.html', 
                           image_data=image_data,
                           total_attempts=total_attempts,
                           total_successful=total_successful,
                           total_failed=total_failed,
                           most_frequent_conversion=most_frequent_conversion,
                           least_frequent_conversion=least_frequent_conversion,
                           never_used_conversion=never_used_conversion,
                           slowest_conversion=slowest_conversion,
                           fastest_conversion=fastest_conversion,
                           conversions_per_minute=conversions_per_minute,
                           conversions_per_hour=conversions_per_hour,
                           conversions_per_day=conversions_per_day)

###To be fetched by the dashboard###

def fetch_document_conversion_data(period):
    # Define the period condition based on the selected value
    if period == 'today':
        period_condition = "DATE(conversion_date) = DATE(NOW())"
    elif period == 'week':
        period_condition = "YEARWEEK(conversion_date, 1) = YEARWEEK(CURDATE(), 1)"
    elif period == 'month':
        period_condition = "MONTH(conversion_date) = MONTH(CURDATE()) AND YEAR(conversion_date) = YEAR(CURDATE())"
    else:
        return {"error": "Invalid period"}

    # Create a connection to the database
    connection = create_connection()
    cursor = connection.cursor(named_tuple=True)

    # Fetch total successful and failed conversions based on the period condition
    cursor.execute(f"""
        SELECT COUNT(*) AS total_successful_conversions 
        FROM doc_conversions 
        WHERE conversion_status = 'successful' AND {period_condition}
    """)
    total_successful = cursor.fetchone().total_successful_conversions

    cursor.execute(f"""
        SELECT COUNT(*) AS total_failed_conversions 
        FROM doc_conversions 
        WHERE conversion_status = 'failed' AND {period_condition}
    """)
    total_failed = cursor.fetchone().total_failed_conversions

    # Fetch the most converted type
    cursor.execute(f""" 
        SELECT conv.conversion_type, COUNT(*) AS count
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful' AND {period_condition}
        GROUP BY conv.conversion_type
        ORDER BY count DESC
        LIMIT 1
    """)
    most_frequent_conversion = cursor.fetchone()

    # Fetch the slowest conversion time
    cursor.execute(f""" 
        SELECT conv.conversion_type, MAX(conv.conversion_time) AS max_conversion_time
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful' AND {period_condition}
        GROUP BY conv.conversion_type
        ORDER BY max_conversion_time DESC
        LIMIT 1
    """)
    slowest_conversion = cursor.fetchone()

    # Fetch the fastest conversion time
    cursor.execute(f""" 
        SELECT conv.conversion_type, MIN(conv.conversion_time) AS min_conversion_time
        FROM doc_conversions conv
        WHERE conv.conversion_status = 'successful' AND {period_condition}
        GROUP BY conv.conversion_type
        ORDER BY min_conversion_time ASC
        LIMIT 1
    """)
    fastest_conversion = cursor.fetchone()

    # Close the cursor and connection
    cursor.close()
    connection.close()

    # Return the conversion statistics as a dictionary
    return {
        'total_successful': total_successful,
        'total_failed': total_failed,
        'most_frequent_conversion': most_frequent_conversion,
        'slowest_conversion': slowest_conversion,
        'fastest_conversion': fastest_conversion
    }

@app.route('/logout')
def logout():
    session.pop('username', None)
    return redirect(url_for('login'))

if __name__ == '__main__':
    #app.run(host='0.0.0.0', port=5000)  # Make sure to listen on all interfaces
    app.run(debug=True)
    #pass
